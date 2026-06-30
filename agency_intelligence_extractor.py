from __future__ import annotations

import csv
import datetime as dt
import hashlib
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document


ROOT = Path.cwd()
RUN_DATE = "2026-05-28"
OUT_BASE = ROOT / f"Agency_Master_Intelligence_Extraction_{RUN_DATE}"


LAYER_SYNONYMS = {
    "Sector": [
        "sector",
        "industry",
        "market",
        "competitor",
        "audience sector",
        "sector intelligence",
        "journey mapping",
    ],
    "Offer": [
        "offer",
        "pricing",
        "packaging",
        "value proposition",
        "customer success",
        "monetization",
        "unit economics",
    ],
    "Marketing": [
        "marketing",
        "campaign",
        "content",
        "seo",
        "aeo",
        "geo",
        "demand",
        "funnel",
        "social media",
        "email",
        "creative",
        "narrative",
        "copywriting",
        "attribution",
    ],
    "Sales": [
        "sales",
        "closing",
        "discovery",
        "objection",
        "qualification",
        "crm",
        "pipeline",
        "follow-up",
        "negotiation",
        "revenue architecture",
    ],
    "Client": [
        "client",
        "onboarding",
        "delivery",
        "retention",
        "expansion",
        "customer success",
        "client journey",
        "client fit",
    ],
    "ClientPartner Acquisition": [
        "clientpartner",
        "client partner",
        "partner acquisition",
        "pre-acquisition",
        "acquisition system",
        "crm system",
    ],
    "Branding": [
        "brand",
        "branding",
        "identity",
        "perception",
        "narrative",
        "visual",
        "symbolic",
        "trust",
        "authority",
        "bois",
    ],
    "Finance": [
        "finance",
        "financial",
        "cash",
        "ledger",
        "treasury",
        "allocation",
        "risk",
        "compliance",
        "tax",
        "runway",
        "profitability",
        "finos",
    ],
    "Automation": [
        "automation",
        "agent",
        "agentic",
        "ai",
        "mcp",
        "api",
        "workflow",
        "runtime",
        "prompt",
        "orchestration",
        "connector",
        "code",
        "plugin",
    ],
    "Operations": [
        "operations",
        "sop",
        "execution",
        "cadence",
        "calendar",
        "delivery architecture",
        "runbook",
        "process",
    ],
    "Management": [
        "management",
        "governance",
        "decision",
        "command",
        "executive",
        "strategy",
        "architecture",
        "vision",
        "blueprint",
    ],
    "Legal": ["legal", "contract", "compliance", "privacy", "terms", "agreement", "risk"],
    "Hiring": ["hiring", "role", "recruit", "team", "talent", "headcount", "department", "operator"],
}


DOMAIN_DEFAULT = {
    "The Sector Drafts": "Sector",
    "Offer Drafts": "Offer",
    "Marketing Drafts": "Marketing",
    "Sales Drafts": "Sales",
    "The Agency Client. Drafts": "Client",
    "The ClientPartner Draft": "ClientPartner Acquisition",
    "Branding Drafts": "Branding",
    "Financial Drafts": "Finance",
    "The Agency Draft 1.2": "Management",
}


LAYER_DEP = {
    "Sector": {
        "depends": "Management, market research, client context",
        "feeds": "Offer, Marketing, Sales, Client, Automation",
        "inputs": "industry data, competitor patterns, audience psychology, market maturity, regulation",
        "outputs": "sector map, opportunity thesis, risk map, language system, execution constraints",
    },
    "Offer": {
        "depends": "Sector, Client, Finance, Sales, Branding",
        "feeds": "Marketing, Sales, Client Delivery, Finance",
        "inputs": "market pain, delivery capability, buyer psychology, pricing constraints, proof",
        "outputs": "offer architecture, packages, pricing logic, risk reducers, value promise",
    },
    "Marketing": {
        "depends": "Sector, Offer, Branding, Finance, Sales feedback",
        "feeds": "Sales, Client acquisition, Revenue forecasting",
        "inputs": "positioning, ICP, offer, creative doctrine, channel data",
        "outputs": "campaigns, content, funnels, leads, attribution signals, demand intelligence",
    },
    "Sales": {
        "depends": "Marketing, Offer, Sector, ClientPartner Acquisition, Branding, Legal",
        "feeds": "Client, Finance, Management, Offer refinement",
        "inputs": "qualified leads, buyer context, offer, CRM data, objections, stakeholder maps",
        "outputs": "closed revenue, pipeline intelligence, win/loss lessons, forecast data",
    },
    "Client": {
        "depends": "Sales promises, Offer, Operations, Legal, Finance",
        "feeds": "Operations, Retention, Expansion, Referrals, Case studies",
        "inputs": "contract, expectations, onboarding data, delivery plan, client goals",
        "outputs": "value realization, retention signals, expansion triggers, delivery intelligence",
    },
    "ClientPartner Acquisition": {
        "depends": "Sector, Offer, Branding, Marketing, Sales",
        "feeds": "Sales, Partnerships, Distribution, Revenue channels",
        "inputs": "partner ICP, power map, audience pools, acquisition paths",
        "outputs": "partner pipeline, acquisition systems, CRM relationships, channel leverage",
    },
    "Branding": {
        "depends": "Sector, Client, Offer, Marketing, Sales, Culture",
        "feeds": "Marketing, Sales, Client experience, Trust and authority",
        "inputs": "identity, sector codes, audience psychology, culture, proof, service reality",
        "outputs": "positioning, narrative, visual language, trust architecture, brand governance",
    },
    "Finance": {
        "depends": "Sales revenue, Client delivery economics, Operations spend, Legal/tax",
        "feeds": "Management, Operations, Hiring, Offer pricing, Risk controls",
        "inputs": "revenue events, expenses, payroll, taxes, budgets, receivables, payables",
        "outputs": "cash position, allocation decisions, risk alerts, profitability reports, forecasts",
    },
    "Automation": {
        "depends": "All layer rules, data contracts, governance, approved tool access",
        "feeds": "All layers",
        "inputs": "workflows, prompts, data, APIs, events, approval rules",
        "outputs": "agents, automations, dashboards, memory updates, execution logs",
    },
    "Operations": {
        "depends": "Client, Offer, Hiring, Automation, Finance",
        "feeds": "Client outcomes, Finance costs, Management controls",
        "inputs": "delivery plan, SOPs, resource capacity, client obligations",
        "outputs": "executed work, quality signals, delivery metrics, capacity needs",
    },
    "Management": {
        "depends": "All layer reporting, Finance, Risk, Governance",
        "feeds": "All layers",
        "inputs": "KPIs, forecasts, risks, decisions, opportunities, constraints",
        "outputs": "priorities, resource allocation, policy, roadmap, escalation decisions",
    },
    "Legal": {
        "depends": "Management intent, Sales terms, Client obligations, Finance compliance",
        "feeds": "All layers",
        "inputs": "claims, contracts, privacy concerns, payment terms, compliance rules",
        "outputs": "approval gates, contract templates, risk decisions, prohibited claims",
    },
    "Hiring": {
        "depends": "Operations demand, Finance capacity, Management priorities",
        "feeds": "Operations and department capability",
        "inputs": "capacity gaps, workload, service roadmap, budget",
        "outputs": "role scorecards, hiring plan, onboarding/training requirements",
    },
}


KPIS = {
    "Sector": "sector fit score, ICP precision, opportunity quality, market-risk confidence, insight-to-action velocity",
    "Offer": "conversion rate, average contract value, gross margin, perceived value ratio, delivery feasibility score, expansion attach rate",
    "Marketing": "CAC, qualified pipeline, ROMI, CTR/CVR, attribution confidence, creative resonance, funnel drop-off, inbound intent volume",
    "Sales": "win rate, sales cycle length, close rate, pipeline velocity, forecast accuracy, deal leakage, objection resolution rate, NRR handoff quality",
    "Client": "time to first value, delivery SLA, client health, retention, NRR, expansion revenue, referral rate, satisfaction",
    "ClientPartner Acquisition": "partner pipeline, acquisition cost, partner-sourced revenue, channel leverage, relationship quality, CRM health",
    "Branding": "trust signal strength, message recall, narrative coherence, authority score, brand-sales alignment, brand-marketing alignment",
    "Finance": "net liquidity, runway days, gross margin, EBITDA, cash conversion cycle, budget variance, LTV:CAC, DSCR, concentration risk",
    "Automation": "workflow uptime, cycle time reduction, error rate, automation coverage, memory freshness, incident count, manual fallback readiness",
    "Operations": "utilization, throughput, SLA adherence, rework rate, delivery quality, capacity load, incident recovery time",
    "Management": "strategic alignment, decision latency, priority completion, risk closure, resource allocation ROI, system health score",
    "Legal": "contract cycle time, compliance incidents, approval SLA, unresolved legal risk, claim substantiation, privacy incidents",
    "Hiring": "time to fill, role readiness, productivity ramp, capacity gap, retention, manager load, training completion",
}


LIFECYCLE = {
    "Sector": "Awareness, Interest, Qualification, Strategy",
    "Offer": "Interest, Engagement, Qualification, Sales, Expansion",
    "Marketing": "Awareness, Interest, Engagement, Qualification",
    "Sales": "Qualification, Sales, Onboarding handoff",
    "Client": "Onboarding, Delivery, Retention, Expansion, Referral",
    "ClientPartner Acquisition": "Awareness, Engagement, Partnership, Referral",
    "Branding": "Awareness, Interest, Trust, Sales assist, Retention",
    "Finance": "Forecasting, Governance, Delivery economics, Expansion protection",
    "Automation": "All lifecycle stages",
    "Operations": "Onboarding, Delivery, Retention",
    "Management": "All lifecycle stages",
    "Legal": "Qualification, Sales, Onboarding, Delivery, Expansion",
    "Hiring": "Delivery, Scaling, Management",
}


TYPE_RULES = [
    ("vision", ["vision"], "Vision doctrine"),
    ("architecture", ["architecture", "architure", "blueprint", "map", "mapping"], "Architecture / system map"),
    ("execution", ["execution", "runbook", "action plan", "daily", "playbook"], "Execution workflow"),
    ("intelligence", ["intelligence", "knowledge", "research", "sector", "psychographic", "competitor"], "Intelligence asset"),
    ("governance", ["governance", "legal", "compliance", "approval", "risk", "policy"], "Governance / risk control"),
    ("metrics", ["metrics", "tracking", "dashboard", "kpi", "roi", "cac", "attribution"], "Measurement / KPI asset"),
    ("prompt", ["prompt", "agent", "ai", "mcp", "api", "connector", "runtime"], "Automation / agent runtime"),
    ("calendar", ["calendar", "cadence", "rhythm"], "Cadence / calendar control"),
    ("sop", ["sop", "procedure", "workflow"], "SOP / procedure"),
    ("financial", ["cash", "financial", "finance", "ledger", "treasury", "tax", "capital"], "Financial control"),
]


RELATIONSHIP_EDGES = [
    (
        "Sector",
        "Offer",
        "Sector intelligence defines economic arena, ICP, language, constraints, category maturity, and proof requirements.",
        "offer-market fit, pricing confidence, conversion rate",
        "Offer becomes generic or mispriced.",
    ),
    (
        "Offer",
        "Marketing",
        "Offer supplies promise, mechanism, value stack, proof, CTA, pricing frame, and risk reducers.",
        "CAC, CVR, qualified pipeline, ROMI",
        "Marketing creates attention that cannot convert.",
    ),
    (
        "Marketing",
        "Sales",
        "Marketing passes qualified demand, source context, lead score, campaign promise, and attribution data.",
        "pipeline velocity, SQL rate, win rate",
        "Sales loses context and wastes follow-up.",
    ),
    (
        "Sales",
        "Client",
        "Sales transfers promises, discovery, stakeholders, risks, success criteria, and contract context.",
        "onboarding speed, retention, NRR",
        "Delivery mismatch, churn, trust loss.",
    ),
    (
        "Client",
        "Operations",
        "Client data becomes delivery plan, SLA, capacity demand, quality feedback, and SOP updates.",
        "SLA, time to value, margin, utilization",
        "Operations works blind and rework increases.",
    ),
    (
        "Operations",
        "Finance",
        "Operations emits cost, capacity, utilization, delivery margin, vendor, payroll, and project economics.",
        "gross margin, runway, budget variance",
        "Finance cannot see profit quality or cash needs.",
    ),
    (
        "Finance",
        "Management",
        "Finance compresses liquidity, risk, profitability, allocation, forecast, and scale capacity into decisions.",
        "runway, EBITDA, decision latency, growth capacity",
        "Management scales into fragility.",
    ),
    (
        "Automation",
        "All Layers",
        "Automation reads rules, events, data contracts, prompts, and approval boundaries; writes logs, dashboards, memory, and actions.",
        "cycle time, uptime, error rate, leverage",
        "Speed increases chaos without governance.",
    ),
    (
        "Legal",
        "All Layers",
        "Legal governs claims, contracts, privacy, compliance, payment terms, IP, risk acceptance, and approval gates.",
        "compliance incidents, approval SLA, contract risk",
        "Revenue promises create unpriced exposure.",
    ),
    (
        "Hiring",
        "Operations",
        "Hiring converts capacity gaps and capability needs into roles, onboarding, training, and accountability.",
        "utilization, SLA, throughput, quality",
        "Delivery bottlenecks constrain revenue.",
    ),
]


def output_dir() -> Path:
    out = OUT_BASE
    i = 1
    while out.exists():
        i += 1
        out = Path(str(OUT_BASE) + f"_{i}")
    out.mkdir(parents=True, exist_ok=True)
    return out


def read_docx(path: Path) -> tuple[str, list[str]]:
    try:
        doc = Document(str(path))
        parts: list[str] = []
        headings: list[str] = []
        for para in doc.paragraphs:
            text = " ".join((para.text or "").split())
            if not text:
                continue
            parts.append(text)
            style = para.style.name if para.style else ""
            if style.lower().startswith("heading") or (
                len(text) < 100
                and (text.isupper() or re.match(r"^[A-Z][A-Za-z0-9 ,:/&()\-.]+:$", text))
            ):
                headings.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [" ".join((cell.text or "").split()) for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts), headings[:12]
    except Exception as exc:  # noqa: BLE001 - extractor should keep going
        return f"DOCX_READ_ERROR: {exc}", []


def read_text(path: Path) -> tuple[str, list[str]]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        return f"TEXT_READ_ERROR: {exc}", []
    headings = [line.strip("# ").strip() for line in text.splitlines() if line.lstrip().startswith("#")][:12]
    return text, headings


def classify_type(name: str, text: str) -> str:
    haystack = f"{name} {text[:3000]}".lower()
    hits: list[str] = []
    for _, keys, desc in TYPE_RULES:
        if any(key in haystack for key in keys):
            hits.append(desc)
    return "; ".join(dict.fromkeys(hits[:3])) if hits else "Reference / source asset"


def classify_layer(rel: str, text: str) -> tuple[str, list[str]]:
    domain = rel.split("\\")[0]
    scores: Counter[str] = Counter()
    haystack = f"{rel} {text[:12000]}".lower()
    for layer, words in LAYER_SYNONYMS.items():
        for word in words:
            scores[layer] += haystack.count(word)
    if domain in DOMAIN_DEFAULT:
        scores[DOMAIN_DEFAULT[domain]] += 8
    rel_low = rel.lower()
    if "finos-plugin" in rel_low:
        scores["Finance"] += 20
        scores["Automation"] += 10
    if "\\bois\\" in rel_low:
        scores["Branding"] += 20
        scores["Automation"] += 8
    if "elite_marketing_agentic_os" in rel_low:
        scores["Marketing"] += 20
        scores["Automation"] += 8
    if "06_ai_operations" in rel_low:
        scores["Sales"] += 12
        scores["Automation"] += 12
    primary = scores.most_common(1)[0][0] if scores else DOMAIN_DEFAULT.get(domain, "Management")
    secondary = [layer for layer, count in scores.most_common(4) if layer != primary and count > 2]
    return primary, secondary


def purpose_for(name: str, dtype: str, layer: str) -> str:
    name_low = name.lower()
    if "vision" in name_low:
        return f"Defines the long-range doctrine and intent for the {layer} layer."
    if "execution" in name_low or "runbook" in name_low or "action plan" in name_low:
        return f"Turns {layer} intelligence into repeatable operating steps."
    if any(word in name_low for word in ["architecture", "architure", "blueprint", "map"]):
        return f"Positions {layer} components, dependencies, controls, and flows inside the agency OS."
    if any(word in name_low for word in ["governance", "legal", "compliance"]):
        return f"Creates control rules that protect {layer} execution from trust, legal, and operational risk."
    if any(word in name_low for word in ["metrics", "tracking", "dashboard"]):
        return f"Defines measurement logic for {layer} performance and decision quality."
    if any(word in name_low for word in ["prompt", "agent", "mcp", "api"]):
        return f"Provides agent/runtime instructions or integration contracts for automating {layer} work."
    if "Reference / source asset" in dtype:
        return f"Preserves reusable {layer} source intelligence, runtime support, or technical implementation detail."
    return f"Contributes reusable {layer} knowledge, context, or operating doctrine."


def automation_potential(ext: str, dtype: str, text: str) -> str:
    haystack = f"{dtype} {text[:4000]}".lower()
    if ext in [".ts", ".py", ".mjs", ".sql", ".yaml", ".yml", ".json", ".jsonl"] or any(
        key in haystack for key in ["agent", "api", "mcp", "automation", "workflow", "runtime", "connector", "event"]
    ):
        return "High"
    if any(key in haystack for key in ["sop", "execution", "calendar", "cadence", "tracking", "dashboard", "crm"]):
        return "Medium-High"
    return "Medium"


def risk_for(layer: str) -> str:
    risks = {
        "Sector": "stale market assumptions, overgeneralized sector models, weak source confidence",
        "Offer": "overpromising, margin leakage, unclear value metric, delivery mismatch",
        "Marketing": "tracking gaps, message-market mismatch, brand inconsistency, channel waste",
        "Sales": "CRM hygiene gaps, unsupported claims, pipeline leakage, discount pressure",
        "Client": "expectation mismatch, poor onboarding, scope creep, weak outcome proof",
        "ClientPartner Acquisition": "partner misalignment, relationship risk, unclear incentives, CRM fragmentation",
        "Branding": "generic output, narrative drift, visual inconsistency, unproven claims",
        "Finance": "liquidity blindness, bad allocation, compliance gaps, financial state outside ledger",
        "Automation": "unsafe autonomy, stale memory, connector failure, insufficient approvals",
        "Operations": "SOP drift, capacity overload, rework, weak incident recovery",
        "Management": "unclear decision rights, priority conflict, stale governance",
        "Legal": "unapproved claims, contract ambiguity, privacy/compliance exposure",
        "Hiring": "underdefined roles, budget-capacity mismatch, training gaps",
    }
    return risks.get(layer, "orphaned ownership, stale versioning, unclear control")


def related_systems(layer: str, secondaries: list[str]) -> str:
    systems = [f"{layer} OS"] + [f"{secondary} OS" for secondary in secondaries[:2]]
    if layer in ["Sales", "Marketing", "ClientPartner Acquisition"]:
        systems += ["CRM", "Revenue Operations"]
    if layer == "Finance":
        systems += ["FinOS", "Ledger", "Treasury", "Dashboard System"]
    if layer == "Branding":
        systems += ["BOIS", "Brand Memory", "Brand Governance"]
    if "Automation" in [layer] + secondaries:
        systems += ["Agent Runtime", "Memory Log", "Workflow Automation"]
    return "; ".join(dict.fromkeys(systems))


def revenue_impact(layer: str) -> str:
    direct = {"Offer", "Marketing", "Sales", "Client", "ClientPartner Acquisition", "Finance"}
    if layer in direct:
        return "Direct: changes acquisition, conversion, retention, expansion, pricing, cash, or forecast quality."
    if layer in ["Branding", "Sector"]:
        return "Indirect-to-direct: improves trust, market selection, messaging, offer fit, and conversion efficiency."
    return "Leverage: improves decision speed, quality, compliance, capacity, or operating margin."


def is_original_workspace_file(path: Path) -> bool:
    if path.name == Path(__file__).name:
        return False
    return not any(part.startswith("Agency_Master_Intelligence_Extraction_") for part in path.parts)


def extract_records() -> list[dict[str, str | int]]:
    records: list[dict[str, str | int]] = []
    for path in sorted(p for p in ROOT.rglob("*") if p.is_file() and is_original_workspace_file(p)):
        rel = str(path.relative_to(ROOT))
        ext = path.suffix.lower() or "[none]"
        if ext == ".docx":
            text, headings = read_docx(path)
        elif ext == ".pyc":
            text = "Generated Python bytecode cache; source .py file carries strategic/technical content."
            headings = []
        elif ext in [
            ".md",
            ".txt",
            ".csv",
            ".json",
            ".jsonl",
            ".yaml",
            ".yml",
            ".ts",
            ".py",
            ".ps1",
            ".sql",
            ".mjs",
            ".example",
            ".gitignore",
            ".gitkeep",
        ]:
            text, headings = read_text(path)
        else:
            text, headings = "", []
        word_count = len(re.findall(r"[A-Za-z][A-Za-z0-9']+", text))
        primary, secondary = classify_layer(rel, text)
        dtype = classify_type(path.stem, text)
        dep = LAYER_DEP.get(primary, LAYER_DEP["Management"])
        records.append(
            {
                "Document Name": path.name,
                "Relative Path": rel,
                "Extension": ext,
                "Word Count": word_count,
                "Purpose": purpose_for(path.stem, dtype, primary),
                "Strategic Function": f"{dtype}; anchors {primary} in the Vision -> Strategy -> System -> Workflow chain.",
                "Business Function": f"Supports {primary} operating performance and its handoffs into {dep['feeds']}.",
                "Layer Assignment": primary + ((" + " + ", ".join(secondary)) if secondary else ""),
                "Dependencies": dep["depends"],
                "Inputs": dep["inputs"],
                "Outputs": dep["outputs"],
                "Workflows Supported": LIFECYCLE.get(primary, "All lifecycle stages"),
                "KPIs Influenced": KPIS.get(primary, "system health, decision quality, revenue efficiency"),
                "Revenue Impact": revenue_impact(primary),
                "Automation Potential": automation_potential(ext, dtype, text),
                "Decision Value": "High"
                if any(key in dtype for key in ["Architecture", "Vision", "Governance", "Measurement", "Automation"])
                or word_count > 2500
                else "Medium",
                "Scalability Value": "High"
                if any(key in dtype for key in ["Architecture", "Execution", "SOP", "Automation", "Governance", "Measurement"])
                else "Medium",
                "Risk Factors": risk_for(primary),
                "Related Documents": f"Same layer documents; upstream: {dep['depends']}; downstream: {dep['feeds']}",
                "Related Systems": related_systems(primary, secondary),
                "Related Departments": "; ".join(dict.fromkeys([primary] + secondary + (["Management"] if primary != "Management" else []))),
                "Lifecycle Position": LIFECYCLE.get(primary, "All lifecycle stages"),
                "Headings / Signals": " | ".join(headings[:5])[:500],
            }
        )
    return records


def write_csv(path: Path, rows: list[dict[str, str | int]]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def duplicate_groups(records: list[dict[str, str | int]]) -> list[list[str]]:
    hash_groups: defaultdict[tuple[str, int], list[str]] = defaultdict(list)
    for record in records:
        rel = str(record["Relative Path"])
        path = ROOT / rel
        ext = str(record["Extension"])
        if ext == ".docx":
            text, _ = read_docx(path)
        elif ext == ".pyc":
            text = ""
        else:
            text, _ = read_text(path)
        digest = hashlib.md5(re.sub(r"\s+", " ", text.lower()).encode("utf-8")).hexdigest()
        hash_groups[(digest, int(record["Word Count"]))].append(rel)
    return [group for (_, words), group in hash_groups.items() if len(group) > 1 and words > 50]


def write_relationship_matrix(out: Path) -> None:
    with (out / "Agency_Cross_System_Relationship_Matrix.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["From Layer", "To Layer", "Information Exchanged", "KPI Impact", "Failure Impact"])
        writer.writerows(RELATIONSHIP_EDGES)


def section_block(
    title: str,
    purpose: str,
    layer: str,
    deps: str,
    inputs: str,
    outputs: str,
    workflow: str,
    auto: str,
    rev: str,
    kpi: str,
    decision: str,
    scale: str,
    risk: str,
    opt: str,
) -> str:
    return f"""
## {title}

| Field | Operating Specification |
|---|---|
| Strategic Purpose | {purpose} |
| Supported Layer | {layer} |
| Dependencies | {deps} |
| Inputs | {inputs} |
| Outputs | {outputs} |
| Workflow Integration | {workflow} |
| Automation Connections | {auto} |
| Revenue Impact | {rev} |
| KPI Impact | {kpi} |
| Decision Value | {decision} |
| Scalability Impact | {scale} |
| Risk Exposure | {risk} |
| Optimization Opportunities | {opt} |
"""


def write_knowledge_graph(out: Path, records: list[dict[str, str | int]]) -> None:
    graph = {
        "vision": "360 degree Growth Revenue Agency Operating System",
        "mission": "Orchestrate client growth across all revenue channels through synchronized domain execution, intelligence-driven decision making, scalable operational systems, and integrated revenue architecture.",
        "topology": ["Sector", "Offer", "Marketing", "Sales", "Client", "Operations", "Finance", "Management"],
        "support_layers": ["Automation", "Legal", "Hiring", "Branding", "ClientPartner Acquisition"],
        "edges": [
            {"from": a, "to": b, "exchange": info, "kpi_impact": kpi, "failure_impact": fail}
            for a, b, info, kpi, fail in RELATIONSHIP_EDGES
        ],
        "documents": [
            {
                "path": record["Relative Path"],
                "layer": record["Layer Assignment"],
                "purpose": record["Purpose"],
                "kpis": record["KPIs Influenced"],
            }
            for record in records
        ],
    }
    (out / "Agency_Knowledge_Graph.json").write_text(json.dumps(graph, indent=2, ensure_ascii=False), encoding="utf-8")


def build_report(out: Path, records: list[dict[str, str | int]], dups: list[list[str]]) -> str:
    file_counts = Counter(str(record["Relative Path"]).split("\\")[0] for record in records)
    word_counts: Counter[str] = Counter()
    layer_counts = Counter(str(record["Layer Assignment"]).split(" + ")[0] for record in records)
    for record in records:
        word_counts[str(record["Relative Path"]).split("\\")[0]] += int(record["Word Count"])

    layer_table = "\n".join(
        f"| {layer} | {LAYER_DEP[layer]['depends']} | {LAYER_DEP[layer]['outputs']} | {LAYER_DEP[layer]['feeds']} | {KPIS[layer]} |"
        for layer in [
            "Sector",
            "Offer",
            "Marketing",
            "Sales",
            "Client",
            "Operations",
            "Finance",
            "Management",
            "Automation",
            "Legal",
            "Hiring",
            "Branding",
            "ClientPartner Acquisition",
        ]
    )
    rel_table = "\n".join(
        f"| {a} -> {b} | {info} | {kpi} | {fail} |" for a, b, info, kpi, fail in RELATIONSHIP_EDGES
    )
    dup_lines = "\n".join("- " + " <-> ".join(group) for group in dups) if dups else "- No exact duplicate content clusters found."

    report = f"""# The Agency - Master Intelligence Extraction and Operating System Blueprint

Generated: {RUN_DATE}  
Workspace: {ROOT}

This pack reconstructs the workspace as one connected 360 degree Growth Revenue Agency Operating System. The extraction read {len(records)} files, including all discovered DOCX, Markdown, CSV, JSON/JSONL, YAML, TypeScript, Python, SQL, PowerShell, and generated cache artifacts. Generated Python bytecode caches are registered as non-source artifacts; their matching source files govern strategic and technical meaning.

## Corpus Map

| Domain Folder | Files | Extracted Words | Operating Role |
|---|---:|---:|---|
"""
    for domain, count in sorted(file_counts.items()):
        role = DOMAIN_DEFAULT.get(domain, "Support")
        report += f"| {domain} | {count} | {word_counts[domain]} | {role} layer and adjacent systems |\n"

    report += f"""

Layer assignment count: {dict(layer_counts)}

## Governing Constitution

Mission: Orchestrate client growth across all revenue channels through synchronized domain execution, intelligence-driven decision making, scalable operational systems, and integrated revenue architecture.

Core principle: every layer serves the layer above and enables the layer below. No layer operates independently. Every workflow must create value upstream and capability downstream.

Primary flow: Sector -> Offer -> Marketing -> Sales -> Client -> Operations -> Finance -> Management.

Support layers: Automation -> all layers; Legal -> all layers; Hiring -> Operations; Branding -> Marketing, Sales, Client, Management; ClientPartner Acquisition -> Sales and Revenue Channels.
"""
    report += section_block(
        "1. Agency Master Architecture Map",
        "Define the agency as a living revenue organism rather than a set of folders.",
        "All layers",
        "The Agency Draft 1.2, domain draft folders, BOIS, FinOS, Marketing Agentic OS, Sales AI Operations",
        "Vision docs, execution maps, source registries, agent cards, code modules, governance docs",
        "A synchronized topology with operating layers, support layers, data loops, and control gates",
        "All domain work routes through the primary flow and support controls",
        "Agent runtimes, retrieval gates, memory logs, event buses, dashboards",
        "Creates compounding revenue by aligning market intelligence, offers, demand, conversion, delivery, cash, and management",
        "North-star revenue, CAC/LTV, runway, conversion, retention, system health",
        "Turns scattered documents into an executive operating map",
        "Enables repeatable client/project deployment",
        "Fragmentation, duplicate systems, stale drafts, unowned decisions",
        "Promote each domain to a controlled OS with source registry, owner, cadence, and dashboard",
    )
    report += f"""
### Layer Dependency Table

| Layer | Depends On | Outputs | Feeds | KPI Stack |
|---|---|---|---|---|
{layer_table}
"""
    report += section_block(
        "2. Complete Layer Dependency Map",
        "Make upstream and downstream dependencies explicit so no layer optimizes locally against the whole system.",
        "All layers",
        "Layer source documents, dependency mapping draft, operating constitution",
        "Sector intelligence, offer economics, campaign data, CRM, delivery metrics, financial events, governance decisions",
        "Dependency graph, failure map, handoff requirements",
        "Every handoff becomes a workflow contract with fields, owners, KPIs, and escalation rules",
        "Automated routing, alerts, handoff packets, memory updates",
        "Prevents revenue leakage caused by broken cross-layer handoffs",
        "Handoff velocity, failure rate, data completeness, owner response time",
        "Shows where a decision must look before execution",
        "Supports modular scaling without chaos",
        "Silent dependency breaks, unowned handoffs, duplicated workflows",
        "Create formal handoff schemas for each edge in the primary flow",
    )
    report += f"""
## 3. Cross-Domain Relationship Matrix

| Relationship | Information / Decisions Exchanged | KPI Impact | Failure Impact |
|---|---|---|---|
{rel_table}
"""
    sections = [
        (
            "4. Revenue Architecture Blueprint",
            "Engineer how revenue is generated, scaled, retained, expanded, forecasted, protected, and automated.",
            "Sector, Offer, Marketing, Sales, Client, Finance, Management",
            "Sector maps, offer architecture, marketing funnels, sales CRM, client delivery, FinOS ledger/event bus",
            "ICP, offer, campaigns, qualified leads, proposals, closed revenue, delivery costs, collections, retention signals",
            "Revenue model, pipeline, forecast, pricing controls, profitability views, expansion loops",
            "Revenue events flow from awareness through sales and delivery into finance and management",
            "CRM automations, lead scoring, sales AI agents, FinOS revenue events, dashboards",
            "Turns value creation into cash and compounds through retention, expansion, referrals, partnerships, and reusable assets",
            "CAC, LTV, LTV:CAC, ACV, gross margin, NRR, win rate, pipeline velocity, runway",
            "Gives leadership a revenue truth model instead of disconnected activity reports",
            "Allows multiple clients/sectors/offers to run through one architecture",
            "Weak attribution, mispriced offers, delivery margin leakage, cash blind spots",
            "Build canonical revenue objects: sector, ICP, offer, campaign, lead, opportunity, client, invoice, profitability unit",
        ),
        (
            "5. Client Lifecycle Architecture Blueprint",
            "Map every asset to awareness through referral so client value compounds after the first sale.",
            "Marketing, Sales, Client, Operations, Finance",
            "Brand trust, marketing demand, sales discovery, offer promise, onboarding and delivery systems",
            "Audience signals, lead scores, CRM stages, contract, onboarding data, delivery outcomes, client health",
            "Lifecycle stages, SLA handoffs, retention and expansion triggers, referral/case-study loops",
            "Lifecycle runs Awareness -> Interest -> Engagement -> Qualification -> Sales -> Onboarding -> Delivery -> Retention -> Expansion -> Referral -> Partnership",
            "CRM, lifecycle automations, client health dashboards, follow-up agents",
            "Protects acquisition cost by increasing retention and expansion",
            "Time to first value, client health, retention, NRR, referral rate, expansion revenue",
            "Shows what each stage needs before it advances",
            "Makes client delivery repeatable across sectors",
            "Sales-delivery mismatch, churn, scope creep, weak proof capture",
            "Create a client lifecycle control board with entry/exit criteria and source-of-truth fields",
        ),
        (
            "6. End-to-End Workflow Architecture",
            "Convert the whole agency into executable daily, weekly, monthly, quarterly, and yearly rhythms.",
            "Operations, Automation, Management",
            "Execution maps, cadence docs, FinOS workflows, Marketing cadence, Sales action plans, BOIS execution flow",
            "Objectives, owners, triggers, data, approvals, KPIs, source evidence",
            "Runbooks, cadence boards, escalation paths, memory logs, dashboards",
            "Daily monitoring; weekly optimization; monthly strategy and finance control; quarterly architecture reset; annual operating model review",
            "Task routing, agent registries, workflow automations, event buses, alerts",
            "Reduces execution drag and rework while increasing throughput",
            "Cycle time, SLA, experiment velocity, decision latency, automation uptime",
            "Turns architecture into action",
            "Supports scale across clients, domains, and operators",
            "Cadence drift, incomplete runbooks, automation without fallback",
            "Standardize every workflow with trigger, owner, input, output, KPI, approval, recovery",
        ),
        (
            "7. Intelligence Architecture Map",
            "Define how knowledge becomes decisions, decisions become workflows, and workflows update knowledge.",
            "Sector, Branding, Marketing, Sales, Finance, Automation",
            "BOIS knowledge graph, Sales paragraph index, Marketing memory, FinOS source intelligence",
            "Documents, paragraphs, evidence anchors, client memory, campaign history, financial events, decision logs",
            "Knowledge graph, retrieval bundles, agent context, evidence-backed recommendations",
            "Observe -> Classify -> Hypothesize -> Test -> Decide -> Encode -> Reuse",
            "Retrieval gates, vector memory, source registries, paragraph indexes, AI memory logs",
            "Improves decision quality and reduces repeated discovery cost",
            "Signal confidence, evidence coverage, decision accuracy, reuse rate",
            "Prevents generic outputs by grounding actions in source context",
            "Knowledge compounds into reusable assets",
            "Stale sources, orphan documents, no version control, over-retrieval without governance",
            "Unify source registries across all domains and attach canonical IDs to every document",
        ),
        (
            "8. Automation Opportunity Matrix",
            "Identify where agentic execution can increase leverage without removing human authority.",
            "Automation plus all layers",
            "Agent registries, MCP/API connector docs, FinOS tools, BOIS runtime, Marketing OS, Sales AI Ops",
            "Events, prompts, APIs, data contracts, approvals, memory objects",
            "Agents, workflows, alerts, dashboards, generated assets, audit logs",
            "Automation routes repeatable work while high-risk decisions escalate to humans",
            "Codex/Claude runtime split, MCP tools, CRM/API connectors, FinOS event bus, BOIS retrieval",
            "Improves throughput, protects margins, reduces leakage, and enables scale",
            "Automation coverage, uptime, error rate, manual fallback readiness, cycle time reduction",
            "Packages repeatable decisions for safe execution",
            "Supports multi-client and multi-domain operations",
            "Unsafe autonomy, stale memory, external actions without approvals",
            "Build automation registry: trigger, action, owner, risk class, data contract, rollback, human gate",
        ),
        (
            "9. Gap Analysis Report",
            "Expose missing pieces required to move from intelligence corpus to operating company.",
            "All layers",
            "All extracted documents and implemented subsystems",
            "Document roles, dependency maps, duplicate clusters, workflow coverage, KPI coverage",
            "Gap backlog ranked by revenue and control impact",
            "Gaps become roadmap workstreams",
            "Dashboard, CRM, source registry, governance, and automation build tasks",
            "Closes revenue leakage and operational chaos before scaling",
            "Gap closure rate, system health, handoff reliability, forecast confidence",
            "Shows what must be built next",
            "Focuses scaling on control before volume",
            "Overbuilding concepts while missing field-level execution controls",
            "Highest priority gaps: global source registry, CRM schema, KPI dictionary, RACI, legal template pack, client onboarding packets, master dashboards, automation approval matrix",
        ),
        (
            "10. Duplication Audit Report",
            "Find duplicate frameworks and version collisions so the OS has canonical sources.",
            "Management, Operations, Knowledge Management",
            "Document hashes, normalized titles, sales registry notes",
            "File paths, text hashes, names, draft numbers, categories",
            "Duplicate clusters, merge/archive/refactor recommendations",
            "Duplicates enter governance review before migration",
            "Source registry flags superseded/active/archive status",
            "Reduces confusion and prevents operators using outdated playbooks",
            "Duplicate closure rate, canonical source coverage",
            "Improves trust in the operating system",
            "Keeps knowledge base lean as corpus grows",
            "Conflicting versions, typo labels, draft collisions",
            "Merge exact duplicates; refactor overlapping conceptual drafts; archive superseded versions after extracting unique content",
        ),
        (
            "11. Governance Architecture",
            "Define decision rights, approvals, risk classes, source lineage, and review cadence.",
            "Management, Legal, Finance, Automation",
            "Sales escalation rules, BOIS governance, FinOS governance, Marketing escalation matrix",
            "Proposed action, evidence, risk class, owner, approval need, rollback",
            "Governance constitution, RACI, approval packets, decision logs, audit trails",
            "All high-risk actions route through approval before execution",
            "Approval routing, immutable logs, risk alerts, memory updates",
            "Protects revenue from legal, trust, finance, and execution risk",
            "Approval SLA, incident count, audit completeness, risk closure",
            "Separates AI recommendation from human-owned authority",
            "Allows more autonomy at low risk while controlling high risk",
            "Unclear ownership, slow approvals, unlogged decisions",
            "Create one agency-wide governance manual with domain-specific appendices",
        ),
        (
            "12. KPI Architecture",
            "Create a measurement spine from market signal to cash and management decision.",
            "All layers",
            "Domain metrics docs, FinOS dashboards, Marketing cadence, Sales tracking, client lifecycle docs",
            "Source data, CRM stages, campaign metrics, delivery data, ledger events, survey/client health",
            "KPI dictionary, dashboard hierarchy, decision thresholds",
            "KPIs trigger daily monitoring, weekly optimization, monthly control, quarterly reset",
            "Dashboards, alerting, attribution models, FinOS reports",
            "Turns growth into measurable, governable compounding",
            "CAC, LTV, win rate, margin, runway, retention, NRR, system health",
            "Enables decision-grade reporting",
            "Prevents scaling blind",
            "Metric sprawl, inconsistent definitions, weak attribution",
            "Build canonical KPI dictionary: formula, owner, cadence, source, threshold, dashboard",
        ),
        (
            "13. System Health Architecture",
            "Define how the agency knows when the organism is healthy, stressed, or failing.",
            "Management, Finance, Operations, Automation",
            "FinOS alerts, Marketing incident rules, Sales governance, BOIS validation",
            "KPIs, incidents, risk signals, workflow uptime, forecast drift, client health",
            "Health dashboard, alert thresholds, recovery playbooks",
            "Daily watchtower and incident response loops",
            "Monitoring, alerts, runbooks, fallback modes",
            "Prevents silent failure that damages revenue or trust",
            "System uptime, incident response time, runway, client health, forecast drift",
            "Compresses complexity into health signals",
            "Supports operator scale",
            "Alert fatigue, missing thresholds, disconnected dashboards",
            "Define red/yellow/green thresholds per layer and require weekly owner review",
        ),
        (
            "14. Scalability Architecture",
            "Show how the agency scales without chaos by making knowledge, workflows, and controls reusable.",
            "Operations, Automation, Hiring, Finance, Management",
            "All OS blueprints, SOPs, agent cards, finance allocation, hiring needs",
            "Canonical docs, templates, workflows, memory, dashboards, role scorecards",
            "Reusable domain OS modules, client pods, training/certification, capacity plans",
            "Each new client/sector runs through the same operating shell with local intelligence",
            "Agent routing, retrieval, SOP automation, dashboards",
            "Increases margin and capacity while preserving quality",
            "Utilization, margin, rework, ramp time, automation coverage",
            "Shows when to hire, automate, or simplify",
            "Prevents founder-dependent execution",
            "Scaling volume before governance, training, finance controls",
            "Create client pod model with domain leads, shared services, FinOS visibility, and QA gates",
        ),
        (
            "15. Risk Architecture",
            "Map strategic, financial, legal, operational, automation, brand, client, and knowledge risks.",
            "Legal, Finance, Management, Automation, Operations",
            "Escalation rules, legal docs, financial risk engine, BOIS governance",
            "Risk signals, claims, contracts, spend, client commitments, data access, automation actions",
            "Risk register, approval paths, incident response, mitigation backlog",
            "Risk gates run before external, financial, legal, destructive, or client-facing action",
            "Risk alerts, approval packets, audit logs, exception reports",
            "Protects revenue, reputation, cash, and client trust",
            "Risk closure rate, compliance incidents, liquidity alerts, churn risk",
            "Makes risk visible before it becomes damage",
            "Supports controlled autonomy and delegated execution",
            "Unpriced risk, undocumented exceptions, unsafe claims",
            "Create agency-wide risk register with owner, severity, trigger, mitigation, review cadence",
        ),
        (
            "16. Knowledge Graph",
            "Connect Vision -> Strategy -> Systems -> Frameworks -> Processes -> Workflows -> Tasks -> Execution -> Data -> Feedback -> Optimization.",
            "All layers",
            "All extracted documents; see Agency_Knowledge_Graph.json",
            "Document metadata, layer assignments, dependencies, KPIs",
            "Machine-readable graph and source-to-layer register",
            "Graph powers retrieval, routing, audit, and roadmap decisions",
            "JSON graph, source registry, future vector retrieval",
            "Turns knowledge into reusable operating memory",
            "Coverage, orphan rate, retrieval precision, source freshness",
            "Shows where every file fits",
            "Prevents orphan assets",
            "Stale graph, missing version status, weak source IDs",
            "Promote graph into a maintained registry with active/superseded/archive flags",
        ),
        (
            "17. Operational Control Framework",
            "Define the control layer that lets the agency run daily without losing strategic coherence.",
            "Operations, Management, Automation",
            "Cadence docs, runbooks, workflow maps, agent routing, finance controls",
            "Objectives, tasks, owners, inputs, outputs, KPIs, approvals, incidents",
            "Control boards, SOPs, escalation queues, decision logs",
            "Daily command, weekly optimization, monthly control, quarterly reset",
            "Agents, dashboards, alerts, memory logs",
            "Improves execution reliability and margin",
            "SLA, throughput, decision latency, rework, incident rate",
            "Converts strategic architecture into operational discipline",
            "Enables multi-operator scale",
            "No owner, unclear trigger, weak fallback",
            "Every workflow must have owner, trigger, SLA, KPI, and recovery procedure",
        ),
        (
            "18. Agency Operating System Blueprint",
            "Package the agency into an executable OS with domain plugins and shared controls.",
            "All layers",
            "Core agency docs plus domain OS folders",
            "Sector, offer, marketing, sales, client, finance, brand, operations, legal, hiring, automation intelligence",
            "A modular agency OS with domain runtimes, source registries, dashboards, governance, and roadmaps",
            "Each client/project invokes required domain modules through a controlled workflow",
            "BOIS, FinOS, Marketing Agentic OS, Sales AI Ops, future Sector/Offer/Client/Operations OS modules",
            "Creates a productized, repeatable, high-margin revenue agency",
            "Revenue growth, operating margin, retention, forecast confidence, system health",
            "Makes the business explainable and executable",
            "Supports franchisable/domain-repeatable expansion",
            "Half-built plugins and inconsistent source governance",
            "Build missing domain plugins: SectorOS, OfferOS, ClientOS, OpsOS, LegalOS, HiringOS",
        ),
        (
            "19. Recommended Folder Structure",
            "Move from draft folders to controlled operating directories without losing lineage.",
            "Management, Knowledge Management, Operations",
            "Current folder map, source registries, duplicate analysis",
            "All documents, code systems, source indexes, runtime assets",
            "Canonical folder structure with active/source/archive/runtime separation",
            "Operators use active OS folders; archives preserve lineage",
            "Retrieval indexes and agents point to canonical active docs",
            "Reduces search time and prevents outdated execution",
            "Canonical source coverage, orphan rate, duplicate closure",
            "Clarifies where truth lives",
            "Scales onboarding and governance",
            "Breaking links, losing source lineage during migration",
            "Recommended root: 00_AGENCY_CONSTITUTION, 01_SECTOR_OS, 02_OFFER_OS, 03_MARKETING_OS, 04_SALES_OS, 05_CLIENT_OS, 06_OPERATIONS_OS, 07_FINANCE_FINOS, 08_BRANDING_BOIS, 09_AUTOMATION_AI_OPS, 10_LEGAL_GOVERNANCE, 11_HIRING_CAPACITY, 12_DASHBOARDS_KPIS, 90_SOURCE_ARCHIVE, 99_EXPERIMENTS",
        ),
        (
            "20. Recommended Future Evolution Roadmap",
            "Prioritize transformation from intelligence corpus to live operating system.",
            "All layers",
            "Gap analysis, subsystem maturity, revenue topology",
            "Current docs, implemented OS assets, missing controls, business goals",
            "Phased roadmap",
            "Phase 1 control; Phase 2 canonicalize; Phase 3 automate; Phase 4 client pods; Phase 5 scale governance",
            "Build registries, dashboards, automations, APIs, memory systems",
            "Compounds revenue by reducing chaos while increasing execution leverage",
            "Roadmap completion, revenue efficiency, system health, retention, margin",
            "Sequences decisions by dependency and risk",
            "Avoids scaling disorder",
            "Building advanced automation before source truth and KPI controls",
            "Next 30 days: source registry + KPI dictionary + CRM schema + governance RACI. Days 31-60: handoff schemas + client onboarding + dashboards. Days 61-90: automate low-risk workflows + run one pilot client pod + update memory loops.",
        ),
    ]
    for args in sections:
        report += section_block(*args)
    report += f"""
## Duplication Findings

{dup_lines}

## Document Register

The full per-document extraction is in `Agency_Document_Intelligence_Register.csv`. It contains one row per discovered file and includes: Document Name, Purpose, Strategic Function, Business Function, Layer Assignment, Dependencies, Inputs, Outputs, Workflows Supported, KPIs Influenced, Revenue Impact, Automation Potential, Decision Value, Scalability Value, Risk Factors, Related Documents, Related Systems, Related Departments, Lifecycle Position, and document signals.

## Output Files

- Agency_Master_Operating_System_Blueprint.md
- Agency_Document_Intelligence_Register.csv
- Agency_Cross_System_Relationship_Matrix.csv
- Agency_Knowledge_Graph.json
- Duplication_and_Gap_Summary.md
"""
    return report


def write_gap_summary(out: Path, dups: list[list[str]]) -> None:
    dup_lines = "\n".join("- " + " <-> ".join(group) for group in dups) if dups else "- No exact duplicate content clusters found."
    gap = f"""# Duplication and Gap Summary

## Exact Duplicate Content Clusters

{dup_lines}

## Critical Missing Systems

1. Global source registry across all domains with source IDs, active/superseded/archive status, owner, cadence, and dependency tags.
2. Agency-wide CRM schema that connects Marketing -> Sales -> Client -> Finance.
3. Canonical KPI dictionary with formula, source, owner, cadence, threshold, and dashboard.
4. Formal RACI for Management, Finance, Legal, Automation, Operations, Hiring, and domain leads.
5. Legal template pack: MSA, SOW, DPA/privacy, claims policy, approval checklist, partner terms.
6. Client onboarding and delivery control pack: intake, success criteria, promise handoff, SLA, health score, expansion/referral triggers.
7. Unified dashboard architecture: executive, revenue, marketing, sales, client, operations, finance, system health.
8. Automation approval matrix and registry: trigger, action, risk class, rollback, manual fallback.
9. Hiring/capacity architecture: role scorecards, pod model, utilization thresholds, training/certification.
10. Cross-domain memory protocol that updates source registries, decision logs, and playbooks after every workflow.

## Recommended Disposition Rules

- Merge exact duplicate documents into one canonical active file, preserve duplicates in archive with pointer to canonical.
- Refactor broad conceptual drafts into controlled doctrine, SOP, template, dashboard, and runtime files.
- Archive typo/collision drafts after extracting unique content.
- Eliminate generated cache artifacts from strategic search indexes.
- Promote implemented systems (BOIS, FinOS, Marketing Agentic OS, Sales AI Ops) into runtime folders with owners and version governance.
"""
    (out / "Duplication_and_Gap_Summary.md").write_text(gap, encoding="utf-8")


def main() -> None:
    out = output_dir()
    records = extract_records()
    write_csv(out / "Agency_Document_Intelligence_Register.csv", records)
    write_relationship_matrix(out)
    write_knowledge_graph(out, records)
    dups = duplicate_groups(records)
    (out / "Agency_Master_Operating_System_Blueprint.md").write_text(build_report(out, records, dups), encoding="utf-8")
    write_gap_summary(out, dups)
    summary = {
        "output_dir": str(out),
        "files_written": sorted(path.name for path in out.iterdir()),
        "records": len(records),
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
    }
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
